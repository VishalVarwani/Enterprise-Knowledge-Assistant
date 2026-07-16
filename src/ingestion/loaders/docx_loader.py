"""
docx_loader.py
--------------
Word document (.docx) loader using python-docx.

Extraction strategy:
  - Paragraphs with heading styles are prefixed with "## " or "### "
    so the chunker can use them as natural split points.
  - Tables are rendered as plain text grids (tab-separated rows).
    This preserves the semantic content of tables without needing
    complex table-aware chunking logic.
  - Empty paragraphs (spacing) are collapsed.
  - Headers/footers are skipped (they repeat on every page and
    add noise to embeddings without retrieval value).

Why python-docx over other approaches:
  - Only mature, maintained library for .docx on pure Python
  - Exposes document structure (paragraphs, styles, tables, runs)
  - textract and mammoth exist but python-docx gives the most control
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Union

from loguru import logger

from .base_loader import BaseLoader, RawDocument

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError("python-docx not installed. Run: pip install python-docx")


class DOCXLoader(BaseLoader):
    """Loads .docx files into RawDocument objects."""

    SUPPORTED_EXTENSIONS = {".docx", ".doc"}

    def can_handle(self, source: Union[str, Path]) -> bool:
        path = Path(source) if not isinstance(source, Path) else source
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def load(self, source: Union[str, Path]) -> list[RawDocument]:
        path = Path(source)
        if not path.exists():
            raise ValueError(f"DOCX file not found: {path}")

        raw_bytes = path.read_bytes()
        file_hash = self._compute_hash(raw_bytes)

        logger.info(f"Loading DOCX: {path.name}")

        try:
            doc = Document(path)
        except Exception as e:
            raise ValueError(f"Failed to open DOCX {path.name}: {e}") from e

        content = self._extract_content(doc)

        if not content.strip():
            logger.warning(f"No extractable text in {path.name}")
            return []

        # Count structural elements for metadata
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)

        metadata = {
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "file_size_bytes": path.stat().st_size,
            "author": self._get_author(doc),
        }

        logger.info(
            f"DOCX loaded: {path.name} | "
            f"{paragraph_count} paragraphs, {table_count} tables"
        )

        return [
            RawDocument(
                content=content,
                source_path=str(path.resolve()),
                source_type="docx",
                name=path.stem,
                metadata=metadata,
                file_hash=file_hash,
            )
        ]

    def load_from_bytes(self, data: bytes, filename: str) -> list[RawDocument]:
        """Load DOCX from bytes (file upload endpoint)."""
        file_hash = self._compute_hash(data)

        try:
            doc = Document(io.BytesIO(data))
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX bytes for {filename}: {e}") from e

        content = self._extract_content(doc)
        if not content.strip():
            return []

        return [
            RawDocument(
                content=content,
                source_path=filename,
                source_type="docx",
                name=Path(filename).stem,
                metadata={
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "file_size_bytes": len(data),
                },
                file_hash=file_hash,
            )
        ]

    def _extract_content(self, doc: "Document") -> str:
        """
        Extract full document content with structure preserved.

        Iterates the document body in order (paragraphs AND tables are
        interleaved in OOXML; python-docx exposes them separately by default
        which breaks logical ordering). We walk the XML directly for correctness.
        """
        parts: list[str] = []

        # Walk body elements in document order
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                paragraph = _element_to_paragraph(doc, element)
                if paragraph:
                    text = self._format_paragraph(paragraph)
                    if text:
                        parts.append(text)

            elif tag == "tbl":
                # Table
                table = _element_to_table(doc, element)
                if table:
                    table_text = self._format_table(table)
                    if table_text:
                        parts.append(table_text)

        return "\n\n".join(parts)

    def _format_paragraph(self, paragraph) -> str:
        """Format a paragraph, prefixing headings with markdown markers."""
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = paragraph.style.name.lower() if paragraph.style else ""

        if "heading 1" in style_name:
            return f"# {text}"
        elif "heading 2" in style_name:
            return f"## {text}"
        elif "heading 3" in style_name or "heading 4" in style_name:
            return f"### {text}"
        elif "list" in style_name or paragraph.paragraph_format.left_indent:
            return f"  • {text}"
        else:
            return text

    def _format_table(self, table) -> str:
        """
        Render a table as plain text (tab-separated rows).

        Why not HTML/Markdown tables:
          - Plain text renders correctly in all downstream contexts
          - Simpler for the chunker to handle consistently
          - Embedding models handle plain text tables better than markdown
        """
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("\t".join(cells))
        return "\n".join(rows)

    def _get_author(self, doc: "Document") -> str:
        """Extract author from document core properties."""
        try:
            props = doc.core_properties
            return props.author or ""
        except Exception:
            return ""


# --- Helpers to walk OOXML body in document order ---
def _element_to_paragraph(doc, element):
    """Wrap an XML paragraph element as a python-docx Paragraph."""
    from docx.text.paragraph import Paragraph
    try:
        return Paragraph(element, doc)
    except Exception:
        return None


def _element_to_table(doc, element):
    """Wrap an XML table element as a python-docx Table."""
    from docx.table import Table
    try:
        return Table(element, doc)
    except Exception:
        return None
