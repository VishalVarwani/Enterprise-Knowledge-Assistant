"""
conftest.py
-----------
Shared pytest fixtures for the EKA test suite.
"""

from __future__ import annotations

from typing import Generator
from unittest.mock import MagicMock

import pytest


def pytest_configure(config):
    config.addinivalue_line("markers", "unit: fast, no external deps")
    config.addinivalue_line("markers", "integration: requires real API keys")
    config.addinivalue_line("markers", "slow: > 5 seconds")
    config.addinivalue_line("markers", "red_team: adversarial security tests")


@pytest.fixture(autouse=True)
def set_test_env(monkeypatch):
    """
    Set required env vars so Settings() doesn't fail in unit tests.
    All values are dummies — no real API calls are made in unit tests.
    """
    monkeypatch.setenv("SUPABASE_URL", "https://test.supabase.co")
    monkeypatch.setenv("SUPABASE_SERVICE_KEY", "test-service-key")
    monkeypatch.setenv("DATABASE_URL", "postgresql://postgres:test@localhost:5432/postgres")
    monkeypatch.setenv("VOYAGE_API_KEY", "test-voyage-key")
    monkeypatch.setenv("GROQ_API_KEY", "test-groq-key")
    # Clear the lru_cache so Settings re-reads env vars per test
    from src.config.settings import get_settings
    get_settings.cache_clear()
    yield
    # Reset PIIFilter singleton so test isolation holds
    try:
        from src.guardrails.pii_filter import PIIFilter
        PIIFilter._analyzer = None
        PIIFilter._anonymizer = None
    except Exception:
        pass


def pytest_collection_modifyitems(config, items):
    if not config.getoption("--run-integration", default=False):
        skip_integration = pytest.mark.skip(reason="Pass --run-integration to run these")
        for item in items:
            if "integration" in item.keywords:
                item.add_marker(skip_integration)


def pytest_addoption(parser):
    parser.addoption(
        "--run-integration",
        action="store_true",
        default=False,
        help="Run integration tests that require real API keys",
    )


# ---------------------------------------------------------------------------
# Sample data fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_text() -> str:
    return (
        "The company's annual leave policy entitles all full-time employees to 25 days "
        "of paid leave per calendar year. Part-time employees receive a pro-rated "
        "entitlement. Leave must be approved by the line manager at least two weeks "
        "in advance, except in cases of emergency."
    )


@pytest.fixture
def sample_pdf_path(tmp_path):
    """Create a PDF with extractable text using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 720), "Leave Policy\nEmployees are entitled to 25 days annual leave per year.")
        pdf_path = tmp_path / "test_policy.pdf"
        doc.save(str(pdf_path))
        doc.close()
        return pdf_path
    except ImportError:
        pytest.skip("PyMuPDF not installed")


@pytest.fixture
def sample_docx_path(tmp_path):
    """Create a minimal DOCX for testing."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Leave Policy", 0)
        doc.add_paragraph("Employees are entitled to 25 days annual leave.")
        doc.add_paragraph("Remote work is allowed up to 3 days per week.")
        path = tmp_path / "test_policy.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture
def sample_chunks():
    """
    Sample RetrievedChunk objects using the actual field names.
    Includes rerank_score as a direct field (set after reranking).
    """
    from src.retrieval.hybrid_search import RetrievedChunk
    return [
        RetrievedChunk(
            chunk_id="chunk-001",
            document_id="doc-001",
            doc_name="HR Policy Manual",
            content="Employees are entitled to 25 days of annual leave per calendar year.",
            doc_source_type="pdf",
            chunk_index=0,
            rrf_score=0.033,
            page_number=5,
            rerank_score=0.95,
            semantic_score=0.91,
            keyword_score=0.80,
        ),
        RetrievedChunk(
            chunk_id="chunk-002",
            document_id="doc-001",
            doc_name="HR Policy Manual",
            content="Leave must be approved by the line manager at least two weeks in advance.",
            doc_source_type="pdf",
            chunk_index=1,
            rrf_score=0.031,
            page_number=5,
            rerank_score=0.88,
            semantic_score=0.84,
            keyword_score=0.72,
        ),
        RetrievedChunk(
            chunk_id="chunk-003",
            document_id="doc-002",
            doc_name="Employee Handbook",
            content="Part-time employees receive a pro-rated leave entitlement.",
            doc_source_type="web",
            chunk_index=12,
            rrf_score=0.029,
            page_number=None,
            rerank_score=0.72,
            semantic_score=0.76,
            keyword_score=0.65,
        ),
    ]


@pytest.fixture
def sample_generation_result():
    from src.generation.generator import GenerationResult
    return GenerationResult(
        answer=(
            "According to the HR Policy Manual [SOURCE 1: HR Policy Manual, p. 5], "
            "employees are entitled to 25 days of annual leave per calendar year. "
            "Leave must be pre-approved [SOURCE 2: HR Policy Manual, p. 5]."
        ),
        citations=["HR Policy Manual, p. 5"],
        grounding_score=0.95,
        has_refusal=False,
        model="llama-3.3-70b-versatile",
        prompt_tokens=450,
        completion_tokens=85,
    )


# ---------------------------------------------------------------------------
# Mock fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def mock_embedder():
    import numpy as np
    embedder = MagicMock()
    embedder.embed_query.return_value = list(np.random.rand(512).astype(float))
    embedder.embed_documents.return_value = [
        list(np.random.rand(512).astype(float)) for _ in range(5)
    ]
    return embedder


@pytest.fixture
def mock_supabase():
    client = MagicMock()
    client.rpc.return_value.execute.return_value.data = []
    client.table.return_value.select.return_value.execute.return_value.data = []
    return client


@pytest.fixture
def mock_groq_response():
    response = MagicMock()
    response.choices = [MagicMock()]
    response.choices[0].message.content = (
        "According to the HR Policy Manual [SOURCE 1: HR Policy Manual, p. 5], "
        "employees receive 25 days of annual leave."
    )
    response.usage.prompt_tokens = 400
    response.usage.completion_tokens = 60
    return response
